import io
import base64
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
from resume_generator import ResumeGenerator

def export_to_pdf(resume_data: dict, template: str = 'modern') -> bytes:
    """Export resume to PDF format"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=12,
            alignment=1,  # Center alignment
            textColor=colors.HexColor('#2c3e50')
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=6,
            spaceBefore=12,
            textColor=colors.HexColor('#34495e')
        )
        
        # Header section
        full_name = resume_data.get('full_name', '')
        contact_info = []
        if resume_data.get('email'):
            contact_info.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_info.append(resume_data['phone'])
        if resume_data.get('location'):
            contact_info.append(resume_data['location'])
        
        elements.append(Paragraph(full_name, title_style))
        elements.append(Paragraph(' | '.join(contact_info), styles['Normal']))
        elements.append(Spacer(1, 12))
        
        # Professional Summary
        if resume_data.get('summary'):
            elements.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
            elements.append(Paragraph(resume_data['summary'], styles['Normal']))
            elements.append(Spacer(1, 12))
        
        # Experience Section
        if resume_data.get('experience'):
            elements.append(Paragraph('PROFESSIONAL EXPERIENCE', heading_style))
            
            for exp in resume_data['experience']:
                # Job title and company
                job_header = f"<b>{exp.get('job_title', '')}</b> - {exp.get('company', '')}"
                elements.append(Paragraph(job_header, styles['Normal']))
                
                # Date and location
                date_range = f"{exp.get('start_date', '')} - {'Present' if exp.get('is_current') else exp.get('end_date', '')}"
                location_date = f"{exp.get('location', '')} | {date_range}"
                elements.append(Paragraph(location_date, styles['Normal']))
                
                # Job description
                if exp.get('description'):
                    desc_lines = exp['description'].split('\n')
                    for line in desc_lines:
                        if line.strip():
                            elements.append(Paragraph(f"• {line.strip()}", styles['Normal']))
                
                elements.append(Spacer(1, 8))
        
        # Education Section
        if resume_data.get('education'):
            elements.append(Paragraph('EDUCATION', heading_style))
            
            for edu in resume_data['education']:
                edu_text = f"<b>{edu.get('degree', '')}</b> in {edu.get('field_of_study', '')}"
                elements.append(Paragraph(edu_text, styles['Normal']))
                
                institution_info = f"{edu.get('institution', '')} | {edu.get('graduation_date', '')}"
                if edu.get('gpa'):
                    institution_info += f" | GPA: {edu['gpa']}"
                elements.append(Paragraph(institution_info, styles['Normal']))
                elements.append(Spacer(1, 8))
        
        # Skills Section
        if resume_data.get('skills'):
            elements.append(Paragraph('SKILLS', heading_style))
            skills_text = ', '.join(resume_data['skills'])
            elements.append(Paragraph(skills_text, styles['Normal']))
        
        # Build PDF
        doc.build(elements)
        
        # Get the value of the BytesIO buffer and return it
        pdf_data = buffer.getvalue()
        buffer.close()
        
        return pdf_data
        
    except Exception as e:
        st.error(f"PDF export failed: {str(e)}")
        return b""

def export_to_docx(resume_data: dict, template: str = 'modern') -> bytes:
    """Export resume to DOCX format"""
    try:
        doc = Document()
        
        # Set up styles
        doc.styles['Normal'].font.name = 'Calibri'
        doc.styles['Normal'].font.size = 11
        
        # Header section
        header = doc.add_heading(resume_data.get('full_name', ''), 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact information
        contact_info = []
        if resume_data.get('email'):
            contact_info.append(resume_data['email'])
        if resume_data.get('phone'):
            contact_info.append(resume_data['phone'])
        if resume_data.get('location'):
            contact_info.append(resume_data['location'])
        
        contact_p = doc.add_paragraph(' | '.join(contact_info))
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Professional Summary
        if resume_data.get('summary'):
            doc.add_heading('PROFESSIONAL SUMMARY', level=1)
            doc.add_paragraph(resume_data['summary'])
        
        # Experience Section
        if resume_data.get('experience'):
            doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
            
            for exp in resume_data['experience']:
                # Job title and company
                job_p = doc.add_paragraph()
                job_run = job_p.add_run(f"{exp.get('job_title', '')} - {exp.get('company', '')}")
                job_run.bold = True
                
                # Date and location
                date_range = f"{exp.get('start_date', '')} - {'Present' if exp.get('is_current') else exp.get('end_date', '')}"
                doc.add_paragraph(f"{exp.get('location', '')} | {date_range}")
                
                # Job description
                if exp.get('description'):
                    desc_lines = exp['description'].split('\n')
                    for line in desc_lines:
                        if line.strip():
                            p = doc.add_paragraph(style='List Bullet')
                            p.add_run(line.strip())
        
        # Education Section
        if resume_data.get('education'):
            doc.add_heading('EDUCATION', level=1)
            
            for edu in resume_data['education']:
                edu_p = doc.add_paragraph()
                edu_run = edu_p.add_run(f"{edu.get('degree', '')} in {edu.get('field_of_study', '')}")
                edu_run.bold = True
                
                institution_info = f"{edu.get('institution', '')} | {edu.get('graduation_date', '')}"
                if edu.get('gpa'):
                    institution_info += f" | GPA: {edu['gpa']}"
                doc.add_paragraph(institution_info)
        
        # Skills Section
        if resume_data.get('skills'):
            doc.add_heading('SKILLS', level=1)
            skills_text = ', '.join(resume_data['skills'])
            doc.add_paragraph(skills_text)
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
        
    except Exception as e:
        st.error(f"DOCX export failed: {str(e)}")
        return b""

def export_to_html(resume_data: dict, template: str = 'modern') -> str:
    """Export resume to HTML format"""
    try:
        generator = ResumeGenerator()
        html_content = generator.generate_html_resume(resume_data, template)
        return html_content
        
    except Exception as e:
        st.error(f"HTML export failed: {str(e)}")
        return "<html><body><h1>Export Failed</h1></body></html>"

def create_download_link(file_data: bytes, filename: str, file_type: str) -> str:
    """Create a download link for file data"""
    try:
        b64 = base64.b64encode(file_data).decode()
        
        mime_types = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'html': 'text/html'
        }
        
        mime_type = mime_types.get(file_type, 'application/octet-stream')
        
        return f'<a href="data:{mime_type};base64,{b64}" download="{filename}">Download {filename}</a>'
        
    except Exception as e:
        st.error(f"Download link creation failed: {str(e)}")
        return ""

def batch_export_resumes(resumes: list, export_format: str) -> bytes:
    """Export multiple resumes in a batch"""
    try:
        import zipfile
        
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for i, resume_data in enumerate(resumes):
                # Generate filename
                name = resume_data.get('full_name', f'Resume_{i+1}').replace(' ', '_')
                filename = f"{name}.{export_format}"
                
                # Export resume
                if export_format == 'pdf':
                    file_data = export_to_pdf(resume_data)
                elif export_format == 'docx':
                    file_data = export_to_docx(resume_data)
                elif export_format == 'html':
                    file_data = export_to_html(resume_data).encode('utf-8')
                else:
                    continue
                
                # Add to zip
                if file_data:
                    zip_file.writestr(filename, file_data)
        
        zip_buffer.seek(0)
        return zip_buffer.getvalue()
        
    except Exception as e:
        st.error(f"Batch export failed: {str(e)}")
        return b""

def validate_export_data(resume_data: dict) -> dict:
    """Validate resume data before export"""
    errors = []
    warnings = []
    
    # Check required fields
    if not resume_data.get('full_name'):
        errors.append("Name is required for export")
    
    if not resume_data.get('email'):
        warnings.append("Email is recommended for professional resumes")
    
    if not resume_data.get('experience'):
        warnings.append("No work experience found")
    
    if not resume_data.get('skills'):
        warnings.append("No skills listed")
    
    return {
        'can_export': len(errors) == 0,
        'errors': errors,
        'warnings': warnings
    }

def get_export_preview(resume_data: dict, export_format: str) -> str:
    """Get a preview of how the exported resume will look"""
    try:
        if export_format == 'html':
            return export_to_html(resume_data)
        elif export_format in ['pdf', 'docx']:
            # Return HTML preview for PDF/DOCX
            return export_to_html(resume_data)
        else:
            return "<p>Preview not available for this format</p>"
            
    except Exception as e:
        return f"<p>Preview failed: {str(e)}</p>"
